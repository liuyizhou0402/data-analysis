# -*- coding: utf-8 -*-
"""Fill the TMGM intern weekly-report template, preserving its original styles."""
import copy
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

SRC = "/root/.claude/uploads/8407763b-0e9d-5767-b77a-e333555d7503/95628876-______.docx"
OUT = "/home/user/data-analysis/tmgm-weekly-report/TMGM_周报_待补充截图.docx"

PH = "【待补充】"   # placeholder marker, highlighted yellow

TASKS = [
    "完成《客户沟通话术与四大资产市场分析》内部培训材料，共 18 页，分为公司介绍话术与市场分析两部分。",
    "梳理 ASIC 产品干预令下的四条合规红线（禁止诱导开户、禁止承诺收益、禁止规避杠杆、禁止个人建议），并整理零售客户杠杆上限对照表（外汇 30:1、黄金与主要股指 20:1、其他商品 10:1、个股 5:1、加密货币 2:1）。",
    "建立 TMGM 事实档案：AFSL 436416 持牌信息、五个集团实体的监管归属区分、以及可对外引用的可查证事实清单，确保对外表述不夸大。",
    "针对新客户、资深交易者、代理 IB 三类对象，分别撰写开场话术、异议应对与合规结尾。",
    "完成黄金、纳斯达克、Alphabet、原油四类资产的基本面与技术面梳理，并补充跨资产联动与估值框架两个专题。",
    "社媒内容运营与客户互动：" + PH + "（本周实际发布与互动情况，数量见下方“社媒情况总结”）",
]

PROBLEMS = [
    "社媒发布的合规边界在实操中不好把握。行情解读与“个人建议”的分界线比较模糊，尤其在私信一对一沟通时，客户经常直接问“现在能不能进”，目前只能回避，但缺少一套既合规又不显得敷衍的标准回应。",
    "行情数据时效性问题。培训材料里的价位、均线数据几天内就会失效，目前没有固定的盘前复核流程，每次对外引用都要重新核一遍，效率低且容易遗漏。",
    "社媒转化链路无法追踪。目前能统计发帖数、评论数、私信数，但从私信到实际开户之间没有记录，判断不出哪类内容真正带来有效线索，只能凭感觉调整内容方向。",
    "点差、佣金、隔夜利息等具体数字必须以官方产品明细表为准，但实习生权限拿不到最新版本，面对资深交易者提问时无法给出确切数值，只能先记录再回复。",
    PH + "（本周实际遇到的其他具体问题，建议补充一条与客户沟通或社媒运营相关的真实案例）",
]

THOUGHTS = [
    "合规限制其实可以讲成差异化卖点。澳洲零售杠杆上限表面上不如离岸平台有吸引力，但把它讲成“保护”而不是“限制”，反而能筛掉赔不起的客户，降低后续的投诉与流失。",
    "内容不该追求预测涨跌，讲框架比讲方向更安全也更有价值。像“机构目标价分歧本身就是波动来源”“指数集中度即风险”这类结构性解释，既不触碰合规红线，又能建立专业形象。",
    "社媒的目标应该是筛选而不是拉量。泛流量带来的多是问“能不能稳赚”的用户，这类客户即使转化，爆仓离开的概率也高。与其追曝光，不如用内容筛出真正理解杠杆风险的人。",
    "需要建立可复用的素材库。合规话术、常见异议应对、市场解读模板如果每次从零开始写，产出不稳定也难以复盘。本周整理的三套话术其实可以直接拆成社媒选题。",
    PH + "（本周个人的其他体会）",
]

doc = Document(SRC)
paras = doc.paragraphs


def write(par, text):
    """Replace a paragraph's text, keeping its first run's formatting."""
    runs = par.runs
    if not runs:
        runs = [par.add_run("")]
    base = runs[0]
    for r in runs[1:]:
        r._element.getparent().remove(r._element)
    if PH in text:
        head, tail = text.split(PH, 1)
        base.text = head
        hl = copy.deepcopy(base._element)
        base._element.addnext(hl)
        from docx.text.run import Run
        hr = Run(hl, par)
        hr.text = PH
        hr.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        hr.font.bold = True
        rest = copy.deepcopy(base._element)
        hl.addnext(rest)
        Run(rest, par).text = tail
    else:
        base.text = text


def clone_after(par):
    new_el = copy.deepcopy(par._element)
    par._element.addnext(new_el)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_el, par._parent)


# --- 本周完成的任务 (slots 2-6) -----------------------------------------
slots = paras[2:7]
while len(slots) < len(TASKS):
    slots.append(clone_after(slots[-1]))
for p, t in zip(slots, TASKS):
    write(p, t)

# --- reload: cloning shifted indices ------------------------------------
doc.save(OUT)
doc = Document(OUT)
paras = doc.paragraphs
idx = {p.text.strip(): i for i, p in enumerate(paras) if p.text.strip()}

def section_slots(header_key, n):
    start = idx[header_key] + 1
    out = []
    i = start
    while len(out) < n and i < len(paras):
        if paras[i].style.name == "List Paragraph" and not paras[i].text.strip():
            out.append(paras[i])
        elif paras[i].text.strip():
            break
        i += 1
    return out

for key, items in [("本周遇到的问题：", PROBLEMS), ("本周的一些思考：", THOUGHTS)]:
    slots = section_slots(key, len(items))
    while len(slots) < len(items):
        slots.append(clone_after(slots[-1]))
    for p, t in zip(slots, items):
        write(p, t)

# --- highlight the social-media count placeholders -----------------------
for p in doc.paragraphs:
    if "xx条" in p.text or "其他平台：" in p.text:
        for r in p.runs:
            if r.text.strip() == "xx" or r.text.strip() == "xx条":
                r.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                r.font.bold = True

# --- note under the title ------------------------------------------------
title = doc.paragraphs[0]
note_el = copy.deepcopy(title._element)
title._element.addnext(note_el)
from docx.text.paragraph import Paragraph
from docx.shared import Pt
note = Paragraph(note_el, title._parent)
for r in note.runs:
    r._element.getparent().remove(r._element)
r = note.add_run("提交前请删除本行：绿色标注处为待你本人填写的内容（社媒数量、个人经历、截图）。")
r.font.size = Pt(10.5)
r.font.bold = True
r.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

doc.save(OUT)
print("saved:", OUT)
