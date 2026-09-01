# -*- coding: utf-8 -*-
"""Broker research sheet - 金融中国, built from the call transcript."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/user/data-analysis/broker-research/券商调研_金融中国.docx"
ASCII_F, CJK_F = "Calibri", "Microsoft YaHei"
INK    = RGBColor(0x1F, 0x2A, 0x37)
ACCENT = RGBColor(0x1F, 0x5E, 0x8C)
MUTED  = RGBColor(0x6B, 0x7A, 0x8C)
FLAG   = RGBColor(0xC0, 0x62, 0x10)   # 待核实
RED    = RGBColor(0xB0, 0x2A, 0x2A)

def font(run, size=None, bold=None, color=None, italic=None):
    if size  is not None: run.font.size = Pt(size)
    if bold  is not None: run.font.bold = bold
    if color is not None: run.font.color.rgb = color
    if italic is not None: run.font.italic = italic
    run.font.name = ASCII_F
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    for k in ('w:ascii', 'w:hAnsi'): rf.set(qn(k), ASCII_F)
    rf.set(qn('w:eastAsia'), CJK_F)
    return run

def shade(cell, hexcolor):
    el = OxmlElement('w:shd'); el.set(qn('w:val'), 'clear'); el.set(qn('w:fill'), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)

def H(doc, text, lvl=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if lvl == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), size=14 if lvl == 1 else 12,
         bold=True, color=INK if lvl == 1 else ACCENT)

def kv_table(doc, rows, w=(4.2, 12.4)):
    """rows: list of (key, value) or (key, value, 'flag')"""
    t = doc.add_table(rows=0, cols=2); t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        k, v = row[0], row[1]
        flag = len(row) > 2 and row[2] == 'flag'
        cells = t.add_row().cells
        for j, (txt, width) in enumerate(zip((k, v), w)):
            cells[j].width = Cm(width)
            par = cells[j].paragraphs[0]
            par.paragraph_format.space_before = Pt(3)
            par.paragraph_format.space_after  = Pt(3)
            font(par.add_run(txt), size=10,
                 bold=(j == 0),
                 color=(FLAG if (flag and j == 1) else INK))
        shade(cells[0], 'EEF3F8')
        if i % 2 == 1: shade(cells[1], 'FAFBFC')
    for r in t.rows:
        for j, c in enumerate(r.cells): c.width = Cm(w[j])
    return t

def note(doc, text, color=MUTED):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    font(p.add_run(text), size=9.5, color=color, italic=True)

doc = Document()
s = doc.sections[0]
s.page_width, s.page_height = Cm(21.0), Cm(29.7)
s.top_margin = s.bottom_margin = Cm(2.0)
s.left_margin = s.right_margin = Cm(2.2)
n = doc.styles['Normal']; n.font.name = ASCII_F; n.font.size = Pt(10.5)
n.element.rPr.rFonts.set(qn('w:eastAsia'), CJK_F)

# ---- title
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
font(p.add_run("券商调研 · 金融中国"), size=20, bold=True, color=INK)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
font(p.add_run("联盟合作政策 · 电话调研纪要"), size=12, color=ACCENT)

kv_table(doc, [
    ("调研对象", "金融中国（香港平台）"),
    ("合作模式称谓", "「联盟」（对方明确表示与市场上的「代理」不同）"),
    ("调研方式", "以潜在代理身份致电官方在线客服转人工，通话约 14 分钟"),
    ("调研日期", "【待填】"),
    ("信息来源", "通话录音转录；对方为一线客服，非专属客户经理"),
    ("信息可靠度", "中。多处数字经语音转录，标橙色处须以官方资料复核"),
])

# ---- 1 佣金
H(doc, "一、联盟佣金结构（每手返佣）")
kv_table(doc, [
    ("第一档（普通用户）", "黄金 2.24 美元 / 手，白银 7.12 美元 / 手", "flag"),
    ("第二档", "黄金约 2.x 美元 / 手，白银约 8.x 美元 / 手", "flag"),
    ("加点权限", "不可加点差、不可加手续费 —— 这是与外部代理模式最大的差别"),
    ("封顶档（VIP8）", "9 美元 / 手"),
])
note(doc, "橙色两行转录严重失真（原文作「2月24」「7月12」「相对于是0的美元」），"
          "数量级可信但小数位必须向对方索取正式佣金表核对。")

# ---- 2 VIP
H(doc, "二、VIP 等级与晋升")
kv_table(doc, [
    ("升级触发条件", "单笔入金 ≥ 20,000 美元，或累计交易 ≥ 1,000 手，系统自动升级"),
    ("VIP 额外回赠", "在联盟佣金之外，每手额外返 2 美元起"),
    ("回赠递增", "VIP1–2 为 2 美元 / 手；VIP3 起升至 4 美元，其后随等级递增 5、6、7、8、9"),
    ("最高等级 VIP8", "每手 9 美元 + 升级赠金 5 万（币种未明） + 每月额外 6 美元", "flag"),
    ("VIP8 晋升门槛", "累计入金 512 万美元，且累计交易 22,000 手"),
])

# ---- 3 点差
H(doc, "三、点差与交易成本")
kv_table(doc, [
    ("点差档位", "仅一档，全体普通用户一致"),
    ("黄金点差", "0.15"),
    ("是否可调", "否。联盟方无法加点，客户之间无差别定价"),
    ("额外费用", "无额外佣金、无额外手续费；出入金仅承担汇率差"),
])

# ---- 4 杠杆
H(doc, "四、杠杆机制（本次调研的关键发现）")
kv_table(doc, [
    ("是否可调", "否。客户与联盟方均无法上调或下调"),
    ("计算方式", "随行情价格动态调整，规则为「价格 ÷ 10」"),
    ("实例", "黄金报价 4,400 美元时，杠杆为 1:440"),
    ("推论", "金价 3,900 时杠杆约 1:390；无法达到量化策略常用的 1:500 或 1:1000"),
])
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6)
font(p.add_run("这是对方产品结构上最实质的短板。"), size=10.5, bold=True, color=RED)
font(p.add_run("杠杆不可调 + 无跟单社区，意味着本地跟单无法做仓位比例缩放，"
               "EA 与量化客户基本无法迁入。团队内部评估认为，代理、工作室或公司类客户"
               "不会选择该平台，主因即在此。"), size=10.5, color=INK)

# ---- 5 跟单
H(doc, "五、跟单、社区与信号源")
kv_table(doc, [
    ("交易软件", "MT5"),
    ("官方跟单社区", "无"),
    ("信号源 / 挂靠服务器", "无"),
    ("本地跟单", "可在 MT5 自行设置，但因杠杆锁定，无法按比例缩放"),
    ("EA / API", "允许。客户手动操作与使用 EA 均被认可"),
    ("刷单", "对方未正面回答", "flag"),
    ("第三方社区", "Wolf 社区（W-O-L-F），内嵌于其 APP。为多家香港交易公司共用的"
                   "公共投资者平台，含群组、好友添加、专家在线讲课"),
])

# ---- 6 结算
H(doc, "六、返佣结算与出入金")
kv_table(doc, [
    ("佣金可见时点", "客户平仓当下即可在收益报表看到"),
    ("提现窗口", "次月 1–8 日，一次性提取上月全部返佣"),
    ("支持币种", "人民币、USDT；港币及其他海外卡走电汇"),
    ("到账速度", "香港本地银行卡出美元一般不走电汇；海外卡电汇较慢"),
    ("手续费", "无额外手续费"),
])

# ---- 7 活动
H(doc, "七、活动政策")
kv_table(doc, [
    ("新客户开户活动", "四个档次，最低 200 美元、中间档 3,000 美元、最高 20,000 美元；"
                       "20,000 档总赠金 800 美元，分三部分发放，并同步升级 VIP"),
    ("中国大陆客户", "可参与"),
    ("老客户活动", "仅偶发的生日活动"),
    ("月度 / 季度活动", "无"),
    ("办公室 / 工作室补贴", "无。对方称客群以个人为主，无面向公司团体的政策"),
    ("下线客户入金物料", "无"),
])
note(doc, "内部评估：活动体系单薄，缺少面向机构型合作方的支持政策，对大体量代理吸引力有限。")

# ---- 8 客服
H(doc, "八、客服与对接体系")
kv_table(doc, [
    ("网页进线", "随机分配客服，无固定归属"),
    ("开户后", "分配专属客服（工号制，本次接触到的工号为 199）"),
    ("入口", "登录交易账户后，页面右上角可见专属客服"),
    ("电话", "默认 AI 应答，发送「转人工」后转接专属客服"),
    ("值班时间", "14:00 – 23:00，非 7×24"),
    ("非值班时段", "由其他坐席协助记录并后续跟进"),
])

# ---- 9 待补
H(doc, "九、待补充 / 待核实事项")
kv_table(doc, [
    ("1", "第一、二档佣金的准确数值（黄金、白银），须索取官方佣金表"),
    ("2", "VIP8 升级赠金「5 万」的币种（人民币或美元）"),
    ("3", "黄金以外品种的点差与佣金：原油、股指、外汇、数字货币"),
    ("4", "VIP1 至 VIP8 的完整晋升门槛表（本次仅拿到 VIP8）"),
    ("5", "是否限制刷单、剥头皮等策略，需书面确认"),
    ("6", "官方物料：佣金表截图、活动页、后台界面截图"),
])

# ---- 10 方法
H(doc, "十、调研方法备忘")
kv_table(doc, [
    ("进线路径", "官网在线客服 → 用任意邮箱注册 → 客服主动对接"),
    ("其他入口", "小红书、抖音等社媒上的官方客服号"),
    ("提效要点", "表明代理身份并强调客户在等，可加快转人工与资料发送"),
    ("资料获取", "由客服在通话过程中于在线聊天窗口发送；对方通常不接受 WhatsApp 传送"),
    ("目标产出", "填写券商调研表，并收集官方图片与物料"),
])

doc.save(OUT)
print("saved:", OUT)
