# -*- coding: utf-8 -*-
"""Build a formatted .docx from the raw transcript text."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/home/user/data-analysis/interview-transcript/transcript_raw.txt"
OUT = "/home/user/data-analysis/interview-transcript/Recording_20260826_0431_逐字转录.docx"

ASCII_FONT = "Calibri"
CJK_FONT   = "Microsoft YaHei"
INK        = RGBColor(0x1F, 0x2A, 0x37)
ACCENT     = RGBColor(0x1F, 0x5E, 0x8C)
MUTED      = RGBColor(0x6B, 0x7A, 0x8C)

# ---------------------------------------------------------------- helpers
def set_font(run, size=None, bold=None, color=None, name=None, italic=None):
    f = run.font
    if size is not None:  f.size = Pt(size)
    if bold is not None:  f.bold = bold
    if italic is not None: f.italic = italic
    if color is not None: f.color.rgb = color
    f.name = name or ASCII_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name or ASCII_FONT)
    rFonts.set(qn('w:hAnsi'), name or ASCII_FONT)
    rFonts.set(qn('w:eastAsia'), CJK_FONT)
    return run

def shade(cell, hexcolor):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear'); el.set(qn('w:fill'), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)

def make_table(doc, rows, widths, header=True):
    t = doc.add_table(rows=0, cols=len(widths))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_i, row in enumerate(rows):
        cells = t.add_row().cells
        for c_i, text in enumerate(row):
            cells[c_i].width = Cm(widths[c_i])
            p = cells[c_i].paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            is_head = header and r_i == 0
            set_font(p.add_run(text), size=9.5, bold=is_head,
                     color=RGBColor(0xFF,0xFF,0xFF) if is_head else INK)
            if is_head:
                shade(cells[c_i], '1F5E8C')
            elif r_i % 2 == 0:
                shade(cells[c_i], 'F4F7FA')
    for row in t.rows:
        for c_i, cell in enumerate(row.cells):
            cell.width = Cm(widths[c_i])
    return t

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(20 if level == 1 else 14)
    pf.space_after  = Pt(8)
    pf.keep_with_next = True
    set_font(p.add_run(text), size=15 if level == 1 else 12.5,
             bold=True, color=INK if level == 1 else ACCENT)
    return p

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.4)
    set_font(p.add_run(text), size=9.5, color=MUTED, italic=True)
    return p

# ---------------------------------------------------------------- document
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin = sec.bottom_margin = Cm(2.2)
sec.left_margin = sec.right_margin = Cm(2.2)

normal = doc.styles['Normal']
normal.font.name = ASCII_FONT
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)

# ---- title block
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
set_font(p.add_run("Recording_20260826_0431"), size=22, bold=True, color=INK)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(16)
set_font(p.add_run("面试录音 · 完整逐字转录"), size=13, color=ACCENT)

make_table(doc, [
    ["项目", "内容"],
    ["录音标识", "Recording_20260826_0431"],
    ["录音时长", "约 54 分 20 秒"],
    ["转录条目", "192 条带时间戳的语音片段"],
    ["转录类型", "原始逐字转录（未润色、未删减）"],
    ["内容概要", "前段为面试前的自我练习与设备调试；29:36 起为正式面试全程"],
], [4.0, 12.6])

heading(doc, "一、内容导航")
note(doc, "以下时间分段为便于检索另行归纳，转录正文本身未作改动。")
make_table(doc, [
    ["时间段", "内容"],
    ["0:03 – 26:26", "面试前自我练习与设备调试：反复演练 “four steps” 答题框架与自我介绍，调整摄像头、画面与视角"],
    ["26:35 – 29:36", "设备测试收尾，确认麦克风与摄像头可用"],
    ["29:36 – 31:08", "面试开始，自我介绍"],
    ["31:08 – 34:52", "案例一 · Johnson & Johnson：皮肤 pre-aging 研究，从约 200 个变量收敛到 9 个关键特征，并构建量化指数"],
    ["34:52 – 35:19", "追问：降维与多变量问题"],
    ["35:21 – 42:13", "案例二 · 创业期客户获取：多渠道试投后聚焦微信群，lead-to-customer 提升约 20%；面试官就“是否只是试错、缺少深入分析”提出质疑"],
    ["42:13 – 47:04", "追问：客户分层逻辑（value / potential / actionability）、时间窗口与潜力的定义方式"],
    ["47:06 – 48:25", "追问：weekly business review 经验与当时的 KPI"],
    ["48:27 – 50:21", "职业方向：business analytics / data engineer / data scientist 的取向"],
    ["50:21 – 50:44", "候选人提问 1：Sydney GBS 团队当前的增长机会与挑战"],
    ["50:44 – 52:41", "面试官作答，并回应候选人提问 2：优秀分析师的共同特质（describe → attribute → predictive → prescriptive 全闭环）"],
    ["52:41 – 54:14", "候选人提问 3：入职 30 / 90 天的成功标准（think-cell → weekly business review → attribution → root cause → 产品熟悉度）"],
    ["54:14 – 54:20", "面试结束"],
], [3.0, 13.6])

heading(doc, "二、关于本转录的说明")
for t in [
    "本文档保留全部时间戳、重复语句、停顿、口误及语音识别可能出现的不准确之处，未作润色或删减。",
    "原始转录将所有语句统一标注为 Speaker1，说话人分离（diarization）未能生效。实际上 29:36 之后的面试部分至少包含两位说话人（面试官与候选人），且说话人切换常出现在同一条记录的中间。正文中省略了逐条重复的 “Speaker1” 标签，其余文字一字未改。",
    "文末附有疑似语音识别错误的对照表，均为根据上下文所作的推测，仅供核对时参考，正文中并未替换。",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    set_font(p.add_run("•  "), size=10.5, color=ACCENT, bold=True)
    set_font(p.add_run(t), size=10.5, color=INK)

# ---- transcript body
heading(doc, "三、逐字转录")

line_re = re.compile(r'^(\d+:\d+)\s+-\s+Speaker(\d+)\s+-\s+(.*)$')
part_re = re.compile(r'^---\s*(Part\s*\d+/\d+)\s*---$')
count = 0
with open(SRC, encoding='utf-8') as fh:
    for raw in fh:
        raw = raw.rstrip('\n').strip()
        if not raw:
            continue
        m = part_re.match(raw)
        if m:
            heading(doc, m.group(1), level=2)
            continue
        m = line_re.match(raw)
        if not m:
            continue
        ts, _spk, text = m.groups()
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(1.9)
        pf.first_line_indent = Cm(-1.9)
        pf.space_after = Pt(5)
        pf.line_spacing = 1.18
        tabs = pf.tab_stops
        tabs.add_tab_stop(Cm(1.9), WD_TAB_ALIGNMENT.LEFT)
        set_font(p.add_run(ts), size=9.5, bold=True, color=ACCENT, name="Consolas")
        p.add_run("\t")
        set_font(p.add_run(text), size=10.5, color=INK)
        count += 1

# ---- appendix
doc.add_page_break()
heading(doc, "附录：疑似语音识别错误对照")
note(doc, "以下为根据上下文推测的可能识别偏差，未在正文中作任何替换。标注“存疑”者为无法从上下文确定的专有名词。")
make_table(doc, [
    ["时间", "转录文字", "疑似原意"],
    ["3:12",  "I will apologize this promise. In eight full. Steps.", "I will approach this problem in four steps."],
    ["6:39",  "TETOL", "公司名称，识别存疑"],
    ["7:37",  "this row", "this role"],
    ["10:02 / 11:44", "stimulation", "simulation"],
    ["10:02 / 30:26", "conversation conversion", "conversion"],
    ["11:47", "my contraction", "my introduction"],
    ["11:59", "I'm Ansel", "I'm Enzo"],
    ["29:58 / 33:45", "filings / filing", "findings"],
    ["30:26", "sales for now", "sales funnel"],
    ["30:26", "Deloitte data platform", "Deloitte data analytics simulation"],
    ["31:44", "preaching", "pre-aging"],
    ["33:45", "sequence secretion", "sebum secretion"],
    ["35:03", "arrivals", "variables"],
    ["35:21", "Florida University", "校名，识别存疑"],
    ["38:20 / 39:20", "red notes / notebook", "小红书（RED）"],
    ["38:45 / 39:20 / 40:08", "term / turnout / turnover", "channel"],
    ["42:45", "action availability", "actionability"],
    ["43:40", "Jess will be contest", "just will be contacted"],
    ["43:40", "low tide", "low tier"],
    ["46:42", "the cosmos", "the customer"],
    ["49:43", "feedback look", "feedback loop"],
    ["49:43", "visit analytics", "business analytics"],
    ["50:44", "group driver", "growth driver"],
    ["53:34", "Once your master get", "Once you master it"],
    ["53:41", "recourse analysis", "root cause analysis"],
    ["54:20", "Hey, Shah", "识别存疑，疑为人名或致谢语"],
], [2.4, 6.6, 7.6])

doc.save(OUT)
print(f"transcript lines written: {count}")
print("saved:", OUT)
